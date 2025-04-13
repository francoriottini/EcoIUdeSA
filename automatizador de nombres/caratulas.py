import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.oxml.ns import qn

def add_instruction(paragraph, parts):
    """Agrega un párrafo con formato personalizado a partir de partes [(texto, subrayado)]"""
    run = None
    for text, is_underlined in parts:
        run = paragraph.add_run(text)
        run.font.name = 'Century Schoolbook'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Century Schoolbook')
        run.font.size = Pt(11)
        run.underline = is_underlined
    paragraph.paragraph_format.space_after = Pt(4)

# Cargar archivo Excel
archivo = r'G:\Otros ordenadores\Mi portátil\UDESA\ECO I\Otoño 2024\Alumnos_Parcial_O2025.xlsx'
df = pd.read_excel(archivo)

# Preparar nombres completos
df = df.dropna(subset=['Apellido', 'Nombre'])
df['Nombre completo'] = df['Apellido'].str.strip() + ', ' + df['Nombre'].str.strip()

# Documento Word
doc = Document()

# Establecer fuente por defecto: Century Schoolbook
style = doc.styles['Normal']
font = style.font
font.name = 'Century Schoolbook'

# Ajustar márgenes
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# Ruta al logo
ruta_logo = r'automatizador de nombres\Figures\Mini-Logo UdeSA.jpg'

for i, nombre in enumerate(df['Nombre completo'], 1):
    # Crear tabla de 1 fila, 2 columnas (logo - nombre)
    encabezado_tabla = doc.add_table(rows=1, cols=2)
    encabezado_tabla.autofit = False
    encabezado_tabla.columns[0].width = Inches(3)
    encabezado_tabla.columns[1].width = Inches(3)

    # Logo en celda izquierda
    cell_logo = encabezado_tabla.cell(0, 0)
    paragraph_logo = cell_logo.paragraphs[0]
    run_logo = paragraph_logo.add_run()
    run_logo.add_picture(ruta_logo, width=Inches(1.5))
    paragraph_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Datos del estudiante en celda derecha
    cell_datos = encabezado_tabla.cell(0, 1)
    paragraph_datos = cell_datos.paragraphs[0]
    paragraph_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Dato del grupo
    grupo = df.loc[i - 1, 'Grupo']

    run_datos = paragraph_datos.add_run(f'Estudiante: {nombre}\nGrupo: {grupo}\nOrden: {i}')
    run_datos.font.size = Pt(11)
    run_datos.bold = True

    # Encabezado general
    p1 = doc.add_paragraph("E010 – Economía I")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.runs[0]
    run1.bold = True
    run1.font.size = Pt(14)

    p2 = doc.add_paragraph("Otoño 2025")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.runs[0]
    run2.font.size = Pt(12)

    p3 = doc.add_paragraph("EXAMEN PARCIAL 1")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.runs[0]
    run3.bold = True
    run3.font.size = Pt(12)

    p4 = doc.add_paragraph("21 de abril")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.runs[0]
    run4.font.size = Pt(12)

    p5 = doc.add_paragraph("Espere a que se le indique cuándo comenzar")
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.runs[0]
    run5.italic = True
    run5.underline = True
    run5.font.size = Pt(11)

    # Instrucciones
    p = doc.add_paragraph()
    run = p.add_run("Descripción e instrucciones")
    run.bold = True
    run.underline = True

    # Instrucciones con marcado para subrayado
    instrucciones_formateadas = [
        [("• La duración del examen es de 130 minutos y consiste en 3 secciones. La primera sección consiste en ", False),
        ("esta parte escrita", True),
        (" y cuenta con ", False),
        ("70 minutos", True),
        (" (1 hora y 10 minutos) para resolverla. Las otras dos secciones (Verdaderos/Falsos y Multiple Choice) se responden en computadora y duran 30 minutos cada una.", False)],
        
        [("• Esta sección (Ejercicios Prácticos) equivale al 36% de la nota final del examen. La sección de Verdaderos/Falsos vale un 31% de esta nota y la de Multiple Choice 33%.", False)],
        
        [("• Cada sección tiene más preguntas que las que tiene que responder, por lo que tiene un cierto grado de elección. ", False),
        ("De esta sección escrita, debe responder solo 2 (dos) ejercicios.", True)],
        
        [("• En los primeros 10 minutos de esta parte del examen les recomendamos leer con cuidado. ", False),
        ("Sólo", True),
        (" durante estos minutos se permite hacer preguntas de aclaración, de forma pública.", False)],
        
        [("• Trate de ser lo más breve, explícito y preciso como sea posible, ya que la claridad de su respuesta contribuirá en parte importante a la nota final. Si decide utilizar gráficos o expresiones matemáticas para hacer un punto, descríbalas a fondo.", False)],
        
        [("• Una vez que termine esta parte del examen, avise al profesor presente para que lo retire. No se puede llevar el texto del examen ni ningún papel borrador.", False)],
        
        [("• ¡Buena suerte!", False)]
    ]

    # Agregar los párrafos al documento
    for parts in instrucciones_formateadas:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_instruction(para, parts)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    # Texto inicial
    run = p.add_run("Marque con una ")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)

    # "cruz" en negrita y subrayado
    run = p.add_run("cruz")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True

    # Continuación
    run = p.add_run(" las preguntas ")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)

    # "elegidas" en negrita y subrayado
    run = p.add_run("elegidas")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True

    # Continuación
    run = p.add_run(" en la columna “")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)

    # "Elección" en negrita y subrayado
    run = p.add_run("Elección")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True

    # Cierre
    run = p.add_run("”:")
    run.font.name = 'Century Schoolbook'
    run.font.size = Pt(11)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Encabezado fusionado
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[2])
    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run("Ejercicios prácticos")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Segunda fila
    row2 = table.rows[1].cells
    for i, texto in enumerate(['N° de ejercicio', 'Elección', 'Nota']):
        p = row2[i].paragraphs[0]
        run = p.add_run(texto)
        run.underline = True

    # Fila 3, 4 y 5
    for idx, row in enumerate(table.rows[2:], start=1):
        p = row.cells[0].paragraphs[0]
        run = p.add_run(str(idx))
        run.bold = True
        row.cells[1].text = ''
        row.cells[2].text = ''

    # Salto de página
    if i < len(df):
        doc.add_page_break()

# Guardar
doc.save('Caratulas_O2025.docx')
print("Documento generado correctamente.")
